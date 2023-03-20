# This Python file uses the following encoding: utf-8
from pathlib import Path

import cv2
import numpy as np

from PySide6.QtWidgets import QMessageBox
from PySide6.QtGui import QIcon

import time
import os
from threading import Timer
import json

from docx import Document
from docx.shared import Inches, Pt


class ExportUtils:
    def __init__(self, parent=None):
        # super().__init__(parent)
        self.exported_dir = "Exports and Snapshots/"
        self.snapshot_dir = "Exports and Snapshots/Snapshots/"
        self.model_dir = "Exports and Snapshots/Model Results/"

    def checkExportDir(self):
            try:
                if (not os.path.exists(self.exported_dir)):
                    os.mkdir(self.exported_dir)
                
                if (not os.path.exists(self.snapshot_dir)):
                    os.mkdir(self.snapshot_dir)
                
                if (not os.path.exists(self.model_dir)):
                    os.mkdir(self.model_dir)
                
                return True
            except:
                return False

    def getCurrentFormattedTime(self): 
        current_time = time.localtime()
        formatted_time = time.strftime("%b-%d-%H-%M-%S", current_time)
        return formatted_time

    def resetTakeSnapshot(self, snapshotButton):
        snapshotButton.setText("Take a Snapshot")
        snapshotButton.setEnabled(True)

    def takeSnapshot(self, snapshotButton, image):
        snapshotButton.setText("Capturing")
        snapshotButton.setEnabled(False)

        formatted_time = self.getCurrentFormattedTime()
        filename = formatted_time + ".png"
        
        if (not self.checkExportDir()):
            print("Could not save snapshot")
            return
        
        try: 
            cv2.imwrite(self.snapshot_dir + filename, cv2.cvtColor(image, cv2.COLOR_RGB2BGR))
            print("Snapshot saved to: " + self.snapshot_dir + filename)

            snapshotButton.setText("Captured")
            Timer(2, self.resetTakeSnapshot, [snapshotButton]).start()
        except:
            print("Could not save snapshot")

    def resetExportText(self, exportCheckBox):
        exportCheckBox.setText("Export Results")
    
    def exportToJSON(self, predictionResult, model_details, cameraEffect):
        try: 
            formatted_time = self.getCurrentFormattedTime()

            json_filename = self.model_dir + formatted_time + '/' + 'results-' + formatted_time + '.json'
            
            json_model_details = {"alt_answer_0": "None"}
            try: 
                if (model_details != ""):
                    json_model_details = {}

                    temp_details = model_details.split('\n')
                    temp_details = temp_details[ : -1]

                    for i in range(len(temp_details)):
                        temp_detail = temp_details[i].split('\t')
                        key = 'alt_answer_' + str(i+1)
                        detail = temp_detail[0] + ' ' + temp_detail[1]

                        json_model_details.update({key: detail})
            except:
                print("Could not extract model details for JSON")
                json_model_details = {"alt_answer_0": "Failed to extract"}


            json_data = {
                'model': predictionResult.model_used,
                'question': predictionResult.question,
                'answer': predictionResult.prediction,
                'details': json_model_details,
                'settings': {
                    'effect': cameraEffect,
                }
            }
            
            with open(json_filename, 'w', encoding='utf-8') as outfile: 
                json.dump(json_data, outfile, ensure_ascii=False, indent=4)
            print('Exported to JSON')
        except: 
            print("Could not export to JSON file")

    def exportResults(self, ui, predictionResult, model_details, cameraEffect, weatherEffects, exportCheckBox):
        print("Exporting...")

        if (not self.checkExportDir()):
            print("Failed to make export directories; could not export.")
            return

        try: 
            document = Document()
            formatted_time = self.getCurrentFormattedTime()

            result_dir = self.model_dir + formatted_time + "/"

            os.mkdir(result_dir)

            self.exportToJSON(predictionResult, model_details, cameraEffect)
            
            doc_name = self.model_dir + formatted_time + "/Results_" + formatted_time + ".docx"

            document.add_heading("Question and Model Prediction")

            answer_paragraph = document.add_paragraph()
            answer_format = answer_paragraph.paragraph_format
            answer_format.left_indent = Inches(0.5)

            modelUsed_run = answer_paragraph.add_run("Model Used: " + predictionResult.model_used)
            modelUsed_font = modelUsed_run.font
            modelUsed_font.name = 'Calibri'
            modelUsed_font.size = Pt(12)

            question_run = answer_paragraph.add_run("\nQuestion Asked: " + predictionResult.question)
            question_font = question_run.font
            question_font.name = 'Calibri'
            question_font.size = Pt(12)

            answer_run = answer_paragraph.add_run("\nPrediction: " + predictionResult.prediction)
            answer_font = answer_run.font
            answer_font.name = 'Calibri'
            answer_font.size = Pt(12)

            if (model_details != ""):
                document.add_heading("Prediction Details")
                details_paragraph = document.add_paragraph()
                details_format = details_paragraph.paragraph_format
                details_format.left_indent = Inches(0.5)

                details_run = details_paragraph.add_run(model_details)
                details_font = details_run.font
                details_font.name = 'Calibri'
                details_font.size = Pt(12)

            document.add_heading("User Settings")
            settings_paragraph = document.add_paragraph()
            settings_format = settings_paragraph.paragraph_format
            settings_format.left_indent = Inches(0.5)

            cameraEffect_run = settings_paragraph.add_run("Camera Effect: " + cameraEffect)
            cameraEffect_font = cameraEffect_run.font
            cameraEffect_font.name = 'Calibri'
            cameraEffect_font.size = Pt(12)

            for weather_condition in weatherEffects: 
                if (weather_condition[1] > 0.0):
                    weather_run = settings_paragraph.add_run("\n" + weather_condition[0] + ": " + str(weather_condition[1]) + "%")
                    weather_font = weather_run.font
                    weather_font.name = 'Calibri'
                    weather_font.size = Pt(12)

            document.add_heading("Base Image")
            cv2.imwrite(result_dir + "base_image.png", cv2.cvtColor(predictionResult.image, cv2.COLOR_RGB2BGR))
            document.add_picture(result_dir + "base_image.png", width=Inches(6.0))

            numVisuals = len(predictionResult.visualizations)
            if numVisuals:
                document.add_heading("Model Visualizations")
                # Export each visual
                for i in range(numVisuals):
                    # Resize Image for Export (720p)
                    dim = (1280, 720)
                    export_frame = cv2.resize(predictionResult.visualizations[i], dim)

                    visual_filename = f"Visualization_{i+1}.png"
                    cv2.imwrite(result_dir + visual_filename, cv2.cvtColor(export_frame, cv2.COLOR_RGB2BGR))
                    
                    document.add_heading(f"Visualization_{i+1}.png", level=2)
                    document.add_picture(result_dir + visual_filename, width=Inches(6.0))
            

            document.save(doc_name)
            print("Results exported to: " + result_dir)

            exportCheckBox.setText("Exported")

            exportSuccessBox = QMessageBox(ui)
            exportSuccessBox.setText("<p>Results exported to <b>" + result_dir + "</b></p>")
            exportSuccessBox.setStyleSheet("""* { background: white; color: #0d0f75; font-family: Arial; font-size: 13px; font-weight: 400; } 
                QPushButton{ background-color: rgb(13, 15, 117); color: white; border-color: blue; border-radius: 5px; padding: 10px 20px; }""")
            exportSuccessBox.setWindowTitle("Export Success")
            exportSuccessBox.setWindowIcon(QIcon("Images/Logos/logo_drone_only.png"))
            exportSuccessBox.show()
            exportSuccessBox.move(exportSuccessBox.pos().x()+10, exportSuccessBox.pos().y())
            exportSuccessBox.exec()

            self.resetExportText(exportCheckBox)
            # Timer(2, self.resetExportText, [exportCheckBox]).start()
            return True
        except:
            exportErrorBox = QMessageBox(ui)
            exportErrorBox.setText("""<p>Error Exporting Results</p>""")
            exportErrorBox.setStyleSheet("""* { background: white; color: #0d0f75; font-family: Arial; font-size: 13px; font-weight: 400; } 
                QPushButton{ background-color: rgb(13, 15, 117); color: white; border-color: blue; border-radius: 5px; padding: 10px 20px; }""")
            exportErrorBox.setWindowTitle("Export Error")
            exportErrorBox.setWindowIcon(QIcon("Images/Logos/logo_drone_only.png"))
            exportErrorBox.show()
            exportErrorBox.move(exportErrorBox.pos().x()+10, exportErrorBox.pos().y())
            exportErrorBox.exec()
            print("Failed to export")
            exportCheckBox.setText("Could Not Export")
            Timer(2, self.resetExportText, [exportCheckBox]).start()
            return False
